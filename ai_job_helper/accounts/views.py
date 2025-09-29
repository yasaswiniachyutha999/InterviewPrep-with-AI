from django.shortcuts import render, redirect
from django.contrib.auth import login
from django.contrib.auth.decorators import login_required
from django.contrib.auth.views import LoginView
from django.contrib.auth.forms import UserCreationForm
from .forms import SignUpForm, UserProfileForm
from .models import UserProfile

# Text extraction helpers
def _extract_text_from_pdf(file_field):
    try:
        from pdfminer.high_level import extract_text
        # Prefer using filesystem path when available
        if hasattr(file_field, 'path') and file_field.path:
            return extract_text(file_field.path)
        # Fallback: open and pass a binary stream
        file_field.open('rb')
        try:
            return extract_text(file_field)
        finally:
            file_field.close()
    except Exception as e:
        print('PDF extract error:', e)
        return ''

def _extract_text_from_docx(file_field):
    try:
        from io import BytesIO
        from docx import Document
        # Use path when available (more reliable)
        if hasattr(file_field, 'path') and file_field.path:
            doc = Document(file_field.path)
        else:
            file_field.open('rb')
            try:
                data = file_field.read()
            finally:
                file_field.close()
            doc = Document(BytesIO(data))
        return '\n'.join(p.text for p in doc.paragraphs)
    except Exception as e:
        print('DOCX extract error:', e)
        return ''

# Import models
from exam.models import Exam
from ats.models import ATSResult
from analysis.models import AnalysisResult

# Custom Login View
class CustomLoginView(LoginView):
    template_name = 'registration/login.html'
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['hide_header'] = True
        return context

# Custom Signup View
def custom_signup(request):
    if request.method == 'POST':
        form = SignUpForm(request.POST)
        if form.is_valid():
            user = form.save()
            login(request, user)
            return redirect('home')
    else:
        form = SignUpForm()
    return render(request, 'registration/signup.html', {'form': form, 'hide_header': True})


def signup(request):
    if request.method == 'POST':
        form = SignUpForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('login')
    else:
        form = SignUpForm()
    return render(request, 'registration/signup.html', {'form': form})


@login_required
def home(request):
    user = request.user

    # Default stats
    stats = {
        "exam_score": "N/A",
        "ats_score": "N/A",
        "resume_score": "N/A",
    }

    # Exam score
    try:
        latest_exam = Exam.objects.filter(user=user).order_by('-created_at').first()
        if latest_exam and latest_exam.score is not None:
            stats["exam_score"] = str(latest_exam.score)  # raw score
    except Exception as e:
        print("Error fetching exam score:", e)

    # ATS score
    try:
        latest_ats = ATSResult.objects.filter(user=user).order_by('-created_at').first()
        if latest_ats and latest_ats.final_score is not None:
            stats["ats_score"] = f"{latest_ats.final_score:.1f}%"
    except Exception as e:
        print("Error fetching ATS score:", e)

    # Resume analysis score
    try:
        latest_analysis = AnalysisResult.objects.filter(user=user).order_by('-created_at').first()
        if latest_analysis and latest_analysis.score is not None:
            stats["resume_score"] = f"{latest_analysis.score:.1f}%"
    except Exception as e:
        print("Error fetching resume score:", e)

    services = [
    {"name": "Resume Builder", "url": "resume_home", "icon": "📄"},
    {"name": "Job Analysis", "url": "analysis_home", "icon": "📊"},
    {"name": "ATS Optimizer", "url": "ats_home", "icon": "⚡"},
    {"name": "Exam Prep", "url": "exam_home", "icon": "📝"},
    {"name": "Training", "url": "training_home", "icon": "📚"},
    {"name": "Mock Interview", "url": "interview_home", "icon": "🎤"},
    {"name": "Portfolio", "url": "portfolio_home", "icon": "💼"},
]

    return render(request, "home.html", {"services": services, "stats": stats})


@login_required
def profile(request):
    profile = request.user.userprofile
    if request.method == 'POST':
        form = UserProfileForm(request.POST, request.FILES, instance=profile)
        if form.is_valid():
            # Save first so storage writes the file and assigns path
            profile = form.save()
            uploaded_file = request.FILES.get('resume_file')
            extracted = ''
            try:
                # Prefer using the uploaded stream immediately if present
                if uploaded_file:
                    name = (uploaded_file.name or '').lower()
                    if name.endswith('.pdf'):
                        extracted = _extract_text_from_pdf(uploaded_file)
                    elif name.endswith('.docx'):
                        extracted = _extract_text_from_docx(uploaded_file)
                # Fallback to stored file on disk
                if not extracted and profile.resume_file:
                    name = (profile.resume_file.name or '').lower()
                    if name.endswith('.pdf'):
                        extracted = _extract_text_from_pdf(profile.resume_file)
                    elif name.endswith('.docx'):
                        extracted = _extract_text_from_docx(profile.resume_file)
            except Exception as e:
                print('Resume extract error:', e)
                extracted = ''
            if extracted:
                profile.extracted_text = extracted
                profile.resume_text = extracted
                profile.save(update_fields=['extracted_text', 'resume_text'])
            return redirect('profile')
    else:
        form = UserProfileForm(instance=profile)
    return render(request, 'accounts/profile.html', {'form': form})
